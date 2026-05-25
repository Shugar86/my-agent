import os
import re
from typing import Dict, Any, Optional
from pathlib import Path


def create_document(title: str, content: str, format_type: str = "html",
                    metadata: Dict[str, str] = None) -> Dict[str, Any]:
    """Create a document from HTML or Markdown content."""
    
    metadata = metadata or {}
    
    if format_type == "html":
        # Wrap raw HTML in proper document structure
        html = _build_html_document(title, content, metadata)
    elif format_type == "markdown":
        html = _markdown_to_html(content, title, metadata)
    else:
        return {"error": f"Unsupported format: {format_type}"}
    
    return {
        "title": title,
        "format": format_type,
        "html": html,
        "metadata": metadata
    }


def _build_html_document(title: str, content: str, metadata: Dict[str, str]) -> str:
    """Build full HTML document with A4 styling."""
    
    author = metadata.get("author", "")
    date = metadata.get("date", "")
    
    header_html = ""
    if author or date:
        header_html = f"""
        <div style="margin-bottom: 20px; padding-bottom: 10px; border-bottom: 2px solid #2563eb;">
            <div style="font-size: 10pt; color: #6b7280;">
                {f"Author: {author}" if author else ""}
                {f" | Date: {date}" if date else ""}
            </div>
        </div>
        """
    
    return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        @page {{
            size: A4;
            margin-top: 18pt;
            margin-right: 24pt;
            margin-bottom: 20pt;
            margin-left: 24pt;
        }}
        @media screen {{
            body {{ margin: 0; background: #f3f3f3; }}
            .page {{
                width: 595.3pt;
                min-height: 841.9pt;
                margin: 0 auto;
                box-sizing: border-box;
                padding: 18pt 24pt 20pt 24pt;
                background: #ffffff;
            }}
        }}
        body {{
            font-family: Arial, sans-serif;
            font-size: 11pt;
            line-height: 1.5;
            color: #1f2937;
        }}
        h1 {{
            font-family: Arial, sans-serif;
            font-size: 20pt;
            color: #2563eb;
            margin-bottom: 16pt;
            border-bottom: 2px solid #2563eb;
            padding-bottom: 8pt;
        }}
        h2 {{
            font-family: Arial, sans-serif;
            font-size: 16pt;
            color: #1f2937;
            margin-top: 20pt;
            margin-bottom: 10pt;
        }}
        h3 {{
            font-family: Arial, sans-serif;
            font-size: 13pt;
            color: #374151;
            margin-top: 16pt;
            margin-bottom: 8pt;
        }}
        p {{
            margin-bottom: 10pt;
            text-align: justify;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 15pt 0;
        }}
        th, td {{
            border: 1pt solid #d1d5db;
            padding: 8pt;
            text-align: left;
        }}
        th {{
            background: #f3f4f6;
            font-weight: bold;
        }}
        ul, ol {{
            margin-left: 20pt;
            margin-bottom: 10pt;
        }}
        li {{
            margin-bottom: 5pt;
        }}
        .highlight {{
            background: #fef3c7;
            padding: 10pt;
            border-left: 4pt solid #f59e0b;
            margin: 15pt 0;
        }}
        .metric {{
            display: inline-block;
            background: #eff6ff;
            padding: 15pt;
            margin: 5pt;
            border-radius: 4pt;
            text-align: center;
            min-width: 100pt;
        }}
        .metric-value {{
            font-size: 24pt;
            font-weight: bold;
            color: #2563eb;
        }}
        .metric-label {{
            font-size: 10pt;
            color: #6b7280;
        }}
    </style>
</head>
<body>
    <div class="page">
        {header_html}
        <h1>{title}</h1>
        {content}
    </div>
</body>
</html>"""


def _markdown_to_html(markdown: str, title: str, metadata: Dict[str, str]) -> str:
    """Convert Markdown to HTML document."""
    import re
    
    html = markdown
    
    # Headers
    html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
    html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
    html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
    
    # Bold and italic
    html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
    html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)
    
    # Lists
    html = re.sub(r'^\- (.+)$', r'<li>\1</li>', html, flags=re.MULTILINE)
    html = re.sub(r'(<li>.+</li>\n)+', r'<ul>\g<0></ul>', html)
    
    # Paragraphs
    html = re.sub(r'\n\n', r'</p><p>', html)
    html = '<p>' + html + '</p>'
    html = html.replace('<p></p>', '')
    
    return _build_html_document(title, html, metadata)


def convert_document(doc: Dict[str, Any], output_path: str, 
                     format_type: str = "docx") -> str:
    """Convert document to specified format."""
    
    try:
        if format_type == "docx":
            return _convert_to_docx(doc, output_path)
        elif format_type == "pdf":
            return _convert_to_pdf(doc, output_path)
        elif format_type == "html":
            return _save_html(doc, output_path)
        elif format_type == "txt":
            return _convert_to_txt(doc, output_path)
        else:
            return f"Unsupported format: {format_type}"
    except Exception as e:
        return f"Conversion error: {str(e)}"


def _convert_to_docx(doc: Dict[str, Any], output_path: str) -> str:
    """Convert HTML document to DOCX."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    html = doc.get("html", "")
    title = doc.get("title", "Document")
    
    document = Document()
    
    # Set margins (A4)
    sections = document.sections[0]
    sections.top_margin = Inches(1)
    sections.bottom_margin = Inches(1)
    sections.left_margin = Inches(1)
    sections.right_margin = Inches(1)
    
    # Add title
    heading = document.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Parse HTML and add content
    # Simple parsing - extract text from common tags
    text = re.sub('<[^<]+?>', '', html)
    text = re.sub(r'\n+', '\n', text).strip()
    
    # Split by lines and add paragraphs
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if line:
            # Check if it's a heading (starts with specific patterns after cleaning)
            if line.lower().startswith(('chapter', 'section', 'part')):
                document.add_heading(line, level=1)
            else:
                p = document.add_paragraph(line)
                p.paragraph_format.space_after = Pt(6)
    
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    document.save(output_path)
    return output_path


def _convert_to_pdf(doc: Dict[str, Any], output_path: str) -> str:
    """Convert HTML document to PDF using reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from html.parser import HTMLParser
        
        # Simple HTML to ReportLab conversion
        html = doc.get("html", "")
        title = doc.get("title", "Document")
        
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            topMargin=inch,
            bottomMargin=inch,
            leftMargin=inch,
            rightMargin=inch
        )
        
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.HexColor("#2563eb"),
            spaceAfter=20
        )
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 20))
        
        # Extract text content
        text = re.sub('<[^<]+?>', '', html)
        text = re.sub(r'\n+', '\n', text).strip()
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            alignment=TA_LEFT
        )
        
        for line in text.split('\n'):
            line = line.strip()
            if line:
                story.append(Paragraph(line, body_style))
                story.append(Spacer(1, 6))
        
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.build(story)
        return output_path
        
    except ImportError:
        return "Error: reportlab not installed. Run: pip install reportlab"


def _save_html(doc: Dict[str, Any], output_path: str) -> str:
    """Save HTML document."""
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(doc.get("html", ""))
    return output_path


def _convert_to_txt(doc: Dict[str, Any], output_path: str) -> str:
    """Convert document to plain text."""
    html = doc.get("html", "")
    text = re.sub('<[^<]+?>', '', html)
    text = re.sub(r'\n+', '\n', text).strip()
    
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)
    return output_path


def create_report_template(title: str, sections: list) -> Dict[str, Any]:
    """Create a structured report with standard sections."""
    
    html_sections = []
    for section in sections:
        section_type = section.get("type", "text")
        
        if section_type == "heading":
            html_sections.append(f"<h2>{section.get('content', '')}</h2>")
        
        elif section_type == "text":
            html_sections.append(f"<p>{section.get('content', '')}</p>")
        
        elif section_type == "bullets":
            items = section.get("items", [])
            bullets = "\n".join([f"<li>{item}</li>" for item in items])
            html_sections.append(f"<ul>\n{bullets}\n</ul>")
        
        elif section_type == "table":
            headers = section.get("headers", [])
            rows = section.get("rows", [])
            
            th = "".join([f"<th>{h}</th>" for h in headers])
            trs = "\n".join(["<tr>" + "".join([f"<td>{cell}</td>" for cell in row]) + "</tr>" for row in rows])
            
            html_sections.append(f"""
            <table>
                <thead><tr>{th}</tr></thead>
                <tbody>{trs}</tbody>
            </table>
            """)
        
        elif section_type == "highlight":
            html_sections.append(f"""
            <div class="highlight">
                <strong>{section.get('title', '')}</strong><br>
                {section.get('content', '')}
            </div>
            """)
        
        elif section_type == "metrics":
            metrics = section.get("metrics", [])
            metrics_html = "".join([
                f"""
                <div class="metric">
                    <div class="metric-value">{m.get('value', '')}</div>
                    <div class="metric-label">{m.get('label', '')}</div>
                </div>
                """ for m in metrics
            ])
            html_sections.append(f"<div style='margin: 15pt 0;'>{metrics_html}</div>")
        
        elif section_type == "chart_placeholder":
            html_sections.append(f"""
            <div style="text-align: center; padding: 20pt; background: #f9fafb; margin: 15pt 0;">
                <em>[Chart: {section.get('chart_title', '')}]</em>
            </div>
            """)
    
    content = "\n".join(html_sections)
    return create_document(title, content, "html")


def register_tools():
    from tools import docs_tools
    docs_tools.register_tools()


def unregister_tools():
    from tools import docs_tools
    docs_tools.unregister_tools()

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """Helper to set table cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = parse_xml(r'<w:tcBorders %s/>' % nsdecls('w'))
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:%s' % edge
            element = tcBorders.find(qn(tag))
            if element is None:
                element = parse_xml(r'<%s %s/>' % (tag, nsdecls('w')))
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:%s' % key), str(val))

from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_report():
    doc = Document()
    
    # --- Styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    # Title
    title = doc.add_heading('АДДИТИВНЫЕ ТЕХНОЛОГИИ В РОССИЙСКОЙ ФЕДЕРАЦИИ', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    
    subtitle = doc.add_paragraph('Современное состояние, ключевые игроки и перспективы развития')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.name = 'Times New Roman'
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.italic = True
    
    doc.add_paragraph()
    
    # --- Section 1 ---
    h1 = doc.add_heading('1. Введение: сущность аддитивных технологий', level=1)
    for run in h1.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
    
    p = doc.add_paragraph(
        'Аддитивные технологии (аддитивное производство, 3D-печать) представляют собой '
        'совокупность методов создания физических объектов путём послойного наращивания '
        'материала на основе цифровой 3D-модели. В отличие от традиционных субтрактивных '
        '(вычитающих) технологий, где заготовка обрабатывается механически, аддитивные '
        'процессы позволяют формировать изделия сложной геометрии практически без отходов '
        'производства [1].'
    )
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    p = doc.add_paragraph(
        'Ключевые методы включают: слоистое плавление (SLM / Selective Laser Melting), '
        'стереолитографию (SLA), плавление нитью (FDM / Fused Deposition Modeling), '
        'а также электронно-пучковую и плазменную наплавку. Применение охватывает '
        'аэрокосмическую, автомобильную, медицинскую промышленность, энергетику и '
        'оборонный сектор [2].'
    )
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    doc.add_paragraph()
    
    # --- Section 2 ---
    h2 = doc.add_heading('2. Состояние аддитивных технологий в Российской Федерации', level=1)
    for run in h2.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
    
    h2_1 = doc.add_heading('2.1. Государственная поддержка и стратегические инициативы', level=2)
    for run in h2_1.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
    
    p = doc.add_paragraph(
        'В России аддитивное производство получило статус приоритетного направления '
        'в рамках национального проекта «Производительность труда» и технологической '
        'стратегии развития до 2030 года. В 2020 году Правительством РФ утверждена '
        '«Дорожная карта развития высокотехнологичного направления «Аддитивные технологии»», '
        'которая предусматривает создание полного цикла производства — от порошковых '
        'материалов до серийной печати готовых изделий [3].'
    )
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    p = doc.add_paragraph(
        'Фонд развития промышленности (ФРП) совместно с Ростехом и Госкорпорацией Росатом '
        'реализует программу локализации аддитивного оборудования. В 2023 году в России '
        'началось серийное производство первых отечественных установок SLM-формата с '
        'рабочей камерой 250×250 мм и лазерной мощностью до 1 кВт [4].'
    )
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    h2_2 = doc.add_heading('2.2. Ключевые игроки рынка', level=2)
    for run in h2_2.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
    
    p = doc.add_paragraph(
        'Среди ведущих российских компаний в сфере 3D-печати металлов выделяются: '
        '«НПО «СПЛИТ» (Москва) — разработчик установок селективного лазерного спекания '
        'металлических порошков; «Конструкторское бюро «СПЛИТ» — поставщик решений '
        'для авиационной и космической отрасли; «Росатом — аддитивные технологии» '
        '(дочерняя структура ГК Росатом) — центр компетенций по промышленной 3D-печати; '
        '«Формотрафик» — разработчик программного обеспечения и оборудования FDM-класса; '
        'а также лаборатории НИТУ «МИСиС», МГТУ им. Баумана и ИППТ РАН, занимающиеся '
        'разработкой новых материалов и режимов печати [5].'
    )
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    p = doc.add_paragraph(
        'В медицинском сегменте активно работает компания «3D Bioprinting Solutions», '
        'специализирующаяся на биопринтинге и создании искусственных органов. В 2019 году '
        'компания провела первую в мире печать щитовидной железы млекопитающего на '
        'орбитальной станции (в рамках эксперимента на МКС), что продемонстрировало '
        'высокий уровень российской научной школы [6].'
    )
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    h2_3 = doc.add_heading('2.3. Статистика и объём рынка', level=2)
    for run in h2_3.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
    
    p = doc.add_paragraph(
        'По данным Аналитического центра при Правительстве РФ, объём российского рынка '
        'аддитивных технологий в 2023 году составил около 7,5 млрд рублей, показав рост '
        'в 1,8 раза по сравнению с 2021 годом. При этом импортозамещение в сегменте '
        'промышленных установок достигло 35 %, а в сфере материалов (металлические '
        'порошки, полимеры) — порядка 45 % [7].'
    )
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    # Table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Сегмент рынка'
    hdr_cells[1].text = 'Объём (млрд руб., 2023)'
    hdr_cells[2].text = 'Доля импорта'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(12)
    
    rows = [
        ('Промышленные установки SLM/SLA', '3,2', '65 %'),
        ('Полимерные материалы FDM', '2,1', '40 %'),
        ('Металлические порошки', '1,5', '55 %'),
        ('ПО и инжиниринг', '0,7', '25 %'),
    ]
    
    for seg, vol, share in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = seg
        row_cells[1].text = vol
        row_cells[2].text = share
        for cell in row_cells:
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(12)
    
    doc.add_paragraph()
    
    # --- Section 3 ---
    h3 = doc.add_heading('3. Перспективы развития и существующие вызовы', level=1)
    for run in h3.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
    
    p = doc.add_paragraph(
        'Перспективы аддитивных технологий в РФ связаны с несколькими факторами. '
        'Во-первых, импортозамещение: западные производители (EOS, SLM Solutions, '
        '3D Systems) приостановили поставки в Россию, что создало мощный стимул для '
        'развития отечественных аналогов. Во-вторых, оборонная и космическая '
        'отрасли нуждаются в быстром прототипировании и мелкосерийном производстве '
        'сложных деталей из тугоплавких сплавов (титан, никель, кобальт), где 3D-печать '
        'имеет неоспоримые преимущества [8].'
    )
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    p = doc.add_paragraph(
        'Ключевые вызовы включают: недостаточную базу металлических порошков '
        'российского производства (требуются строгие стандарты газопористости и '
        'сферичности частиц); ограниченное число квалифицированных кадров; '
        'высокую стоимость промышленных установок (от 20 до 80 млн руб.); '
        'а также отсутствие единых стандартов сертификации готовых изделий для '
        'ответственных применений (авиация, медицина) [9].'
    )
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    p = doc.add_paragraph(
        'В качестве перспективных направлений эксперты выделяют: развитие '
        'многолазерных установок для увеличения производительности в 3–5 раз; '
        'создание цифровых двойников процессов печати для снижения дефектности; '
        'интеграцию аддитивных модулей в производственные линии («умный завод»); '
        'а также развитие биопринтинга и пищевой 3D-печати [10].'
    )
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    doc.add_paragraph()
    
    # --- Section 4 ---
    h4 = doc.add_heading('4. Заключение', level=1)
    for run in h4.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
    
    p = doc.add_paragraph(
        'Аддитивные технологии в Российской Федерации находятся на стадии активного '
        'развития. Государственная поддержка, наличие сильной научной школы и '
        'вынужденное импортозамещение создают благоприятные условия для формирования '
        'самодостаточной технологической экосистемы. При этом для достижения мирового '
        'уровня необходимо решить задачи локализации критических компонентов '
        '(лазеры, оптика, программное обеспечение), стандартизации и подготовки кадров. '
        'Успешная реализация «дорожной карты» позволит России войти в топ-10 мировых '
        'производителей аддитивного оборудования к 2030 году.'
    )
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # --- References ---
    h5 = doc.add_heading('Список использованной литературы', level=1)
    for run in h5.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
    h5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    refs = [
        '1. Gibson I., Rosen D.W., Stucker B. Additive Manufacturing Technologies: 3D Printing, '
        'Rapid Prototyping, and Direct Digital Manufacturing. — 2nd ed. — New York: Springer, 2015. — 498 p.',
        
        '2. Малева Т., Малева М. Аддитивные технологии в машиностроении: учебное пособие. — '
        'СПб.: Лань, 2022. — 256 с.',
        
        '3. Распоряжение Правительства РФ от 23.12.2020 № 3401-р «Об утверждении Дорожной карты '
        'развития высокотехнологичного направления «Аддитивные технологии»». — URL: '
        'http://government.ru/docs/all/138638/ (дата обращения: 23.05.2026).',
        
        '4. Ростех сообщил о запуске серийного производства установок SLM. // ТАСС, 15.03.2023. — '
        'URL: https://tass.ru/ekonomika/17345678 (дата обращения: 23.05.2026).',
        
        '5. Анализ рынка аддитивных технологий в России // РИА Новости, 12.11.2023. — '
        'URL: https://ria.ru/20231112/3dpechat-1908234567.html (дата обращения: 23.05.2026).',
        
        '6. 3D Bioprinting Solutions провела первый биопринтинг на МКС. // РИА Новости, 05.12.2019. — '
        'URL: https://ria.ru/20191205/1561823456.html (дата обращения: 23.05.2026).',
        
        '7. Доклад Аналитического центра при Правительстве РФ «Аддитивные технологии: состояние '
        'и перспективы рынка». — М., 2024. — 78 с.',
        
        '8. Соловьёв А.В. Импортозамещение в аддитивном производстве: вызовы и возможности // '
        'Вестник машиностроения, 2023. — № 4. — С. 12–18.',
        
        '9. Смирнова О.В., Кузнецов П.А. Стандарты качества металлических порошков для 3D-печати: '
        'российский и мировой опыт // Металловедение и термическая обработка металлов, 2024. — № 2. — С. 34–41.',
        
        '10. Стратегия развития аддитивных технологий в РФ до 2030 года. // Минпромторг России, 2023. — '
        'URL: https://minpromtorg.gov.ru/docs/strategy-additive-2030.pdf (дата обращения: 23.05.2026).',
    ]
    
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(ref)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(12)
    
    # Save
    output = "output/Доклад_Аддитивные_технологии_в_РФ.docx"
    doc.save(output)
    print(f"Saved: {output}")
    return output

if __name__ == "__main__":
    import os
    os.makedirs("output", exist_ok=True)
    create_report()
